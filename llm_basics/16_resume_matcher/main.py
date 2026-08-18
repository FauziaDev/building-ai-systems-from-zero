import os
import json
import time
from pathlib import Path

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader
from docx import Document
# 1. ENVIRONMENT + GROQ CLIENT
load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError(
        "GROQ_API_KEY nahi mila. .env file mein GROQ_API_KEY set karo."
    )

client = Groq(api_key=my_api_key)

MODEL = "openai/gpt-oss-120b"
# 2. JOB DESCRIPTION
job_description = """
Description

Do you want to solve real customer problems through innovative technology?
Do you enjoy working on scalable services in a collaborative team environment?
Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf
of our customers. Customer obsession is part of our company DNA, which has
made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve
complex problems while seeing their work's impact first-hand.

We seek individuals passionate about creating new products, features, and
services while managing ambiguity in an environment where development cycles
are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the
entire lifecycle of your code - from design through deployment and ongoing
operations.

Key job responsibilities

• Collaborate and communicate effectively with experienced cross-disciplinary
Amazonians to design, build, and operate innovative products and services.

• Design and develop scalable solutions using cloud-native architectures and
microservices in a large distributed computing environment.

• Participate in code reviews and contribute to technical documentation.

• Build and maintain resilient distributed systems that are scalable,
fault-tolerant, and cost-effective.

• Leverage and contribute to the development of GenAI and AI-powered tools
to enhance development productivity.

• Write clean, maintainable code following best practices and design patterns.

• Work in an agile environment practicing CI/CD principles.

• Demonstrate operational excellence through monitoring, troubleshooting,
and resolving production issues.

Basic Qualifications

- Experience with at least one general-purpose programming language such as
Java, Python, C++, C#, Go, Rust, or TypeScript

- Experience with data structure implementation, basic algorithm development,
and/or object-oriented design principles

- Currently has, or is in the process of obtaining a bachelor's degree in
Computer Science, Computer Engineering, Data Science, Information Systems,
or related STEM fields

- Must be 18 years of age or older

Preferred Qualifications

- Experience from previous technical internship(s) or demonstrated project
experience

- Experience with AI tools for development productivity

- Cloud platforms, preferably AWS

- Database systems, SQL and NoSQL

- Contributing to open-source projects

- Version control systems

- Debugging and troubleshooting complex systems

- Demonstrated ability to learn and adapt to new technologies quickly

- Basic understanding of software development lifecycle (SDLC)

- Strong problem-solving and analytical skills

- Excellent written and verbal communication skills
"""
# 3. PYDANTIC MODELS
class JobD(BaseModel):
    role: str
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    minimum_experience: float | None = None
    education_requirements: list[str] = Field(default_factory=list)
    responsibilities: list[str] = Field(default_factory=list)


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None

    total_experience_years: float | None = None

    skills: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    education: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)


class MatchDetails(BaseModel):
    candidate_name: str | None = None
    matching_skills: list[str] = Field(default_factory=list)
    missing_important_skills: list[str] = Field(default_factory=list)
    experience_requirement_met: bool | None = None
    education_match: bool | None = None
    final_verdict: str = ""


class MatchResult(BaseModel):
    score: float = Field(ge=0, le=100)
    details: MatchDetails
# 4. HELPER FUNCTION - GROQ CALL
def call_llm(messages, retries=4):
    """
    Groq API ko safely call karta hai.

    Temporary errors / rate limits ke case mein retry karega.
    """

    for attempt in range(retries):

        try:

            response = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                response_format={"type": "json_object"},
                temperature=0,
            )

            content = response.choices[0].message.content

            if not content:
                raise ValueError(
                    "LLM ne empty response diya."
                )

            return content

        except Exception as e:

            print(
                f"LLM error "
                f"(attempt {attempt + 1}/{retries}): {e}"
            )

            if attempt < retries - 1:

                # Increasing delay for rate-limit errors
                wait_time = 3 * (attempt + 1)

                print(
                    f"Retrying after {wait_time} seconds..."
                )

                time.sleep(wait_time)

            else:

                raise
# 5. PARSE JOB DESCRIPTION
def parse_job_description(job_text):

    job_schema = JobD.model_json_schema()

    system_prompt = f"""
You are an expert HR assistant.

Analyze the job description and extract structured information.

Return ONLY valid JSON.

The JSON must contain actual extracted values.

Do NOT return the schema itself.

Schema:

{json.dumps(job_schema, indent=2)}

Rules:

1. Do not invent information.

2. If minimum experience is not mentioned,
   return null.

3. If a list has no information,
   return [].

4. Extract required and preferred skills separately.

5. Keep the role concise.

6. Extract important responsibilities.

7. Understand phrases such as:
   "at least one programming language"
   "such as Java, Python, C++, C#, Go, Rust, or TypeScript"
   as an OR requirement, not an AND requirement.

8. Do not interpret alternative technologies
   as individually mandatory requirements.
"""

    user_prompt = f"""
Analyze this job description:

{job_text}
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt,
        },
        {
            "role": "user",
            "content": user_prompt,
        },
    ]

    raw_output = call_llm(messages)

    try:

        data = json.loads(raw_output)

        return JobD(**data)

    except (
        json.JSONDecodeError,
        ValidationError
    ) as e:

        print("Job parsing error:", e)

        print("Raw output:")
        print(raw_output)

        raise
# 6. READ PDF
def read_pdf(file_path):

    try:

        reader = PdfReader(
            str(file_path)
        )

        text_parts = []

        for page_number, page in enumerate(
            reader.pages,
            start=1
        ):

            try:

                page_text = page.extract_text()

                if page_text:

                    text_parts.append(
                        page_text
                    )

            except Exception as e:

                print(
                    f"Warning: PDF page "
                    f"{page_number} read nahi ho paya: {e}"
                )

        return "\n".join(text_parts)

    except Exception as e:

        print(
            f"PDF read error: {e}"
        )

        return None
# 7. READ DOCX
def read_docx(file_path):

    try:

        document = Document(
            str(file_path)
        )

        text_parts = []

        # Paragraphs
        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text:

                text_parts.append(
                    text
                )

        # Tables
        for table in document.tables:

            for row in table.rows:

                row_text = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:

                        row_text.append(
                            cell_text
                        )

                if row_text:

                    text_parts.append(
                        " | ".join(row_text)
                    )

        return "\n".join(text_parts)

    except Exception as e:

        print(
            f"DOCX read error: {e}"
        )

        return None

# 8. READ RESUME
def read_resume(file_path):

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":

        return read_pdf(file_path)

    elif suffix == ".docx":

        return read_docx(file_path)

    else:

        return None
# 9. PARSE RESUME
def parse_resume(resume_text):

    resume_schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume based on meaning,
not only exact section headings.

For example:

Experience
Professional Experience
Work History
Employment
Internships

may all contain work experience.

Skills may appear in:

- Skills
- Experience
- Internships
- Projects
- Certifications

Return ONLY valid JSON.

Schema:

{json.dumps(resume_schema, indent=2)}

Rules:

1. Do not invent information.

2. Missing single values should be null.

3. Missing lists should be [].

4. Include internships inside experiences.

5. Extract skills from the entire resume.

6. Extract projects separately.

7. Extract education information.

8. total_experience_years should be numeric
   when clearly available.

9. Preserve the actual skills found in the resume.

10. Recognize common equivalent terms.

For example:

OOP = OOPS = Object-Oriented Programming

DSA = Data Structures and Algorithms

GitHub is evidence of Git/version control usage
when clearly demonstrated.

Machine Learning and AI projects should be
captured as technical project experience.
"""

    user_prompt = f"""
Parse the following resume:

{resume_text}
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt,
        },
        {
            "role": "user",
            "content": user_prompt,
        },
    ]

    raw_output = call_llm(messages)

    try:

        data = json.loads(raw_output)

        return Resume(**data)

    except (
        json.JSONDecodeError,
        ValidationError
    ) as e:

        print(
            "Resume parsing error:",
            e
        )

        print(
            "Raw output:"
        )

        print(
            raw_output
        )

        raise
# 10. FINAL SCORE
def final_score(job, resume):

    match_schema = MatchResult.model_json_schema()

    prompt = f"""
You are an expert technical recruiter evaluating a candidate for an
entry-level Software Development Engineer (SDE-I) position.

Compare the candidate's resume against the job requirements.

JOB INFORMATION:

{job.model_dump_json(indent=2)}


CANDIDATE RESUME:

{resume.model_dump_json(indent=2)}


Return ONLY valid JSON matching this schema:

{json.dumps(match_schema, indent=2)}
IMPORTANT EVALUATION RULES
1. SCORE RANGE
The score must be between 0 and 100.

Use approximately this weighting:

Core required qualifications       50%
Education                          15%
Internship / project experience   10%
Preferred technical skills        15%
Software engineering evidence     10%

2. PROGRAMMING LANGUAGE RULE
THIS IS VERY IMPORTANT.

If the job says:

"Experience with at least one general-purpose programming
language such as Java, Python, C++, C#, Go, Rust, or TypeScript"

then these languages are ALTERNATIVES.

They are NOT all individually required.

For example, if the candidate has:

Python

then the programming-language requirement is satisfied.

If the candidate has:

Python + C++

then it is also satisfied.

DO NOT penalize the candidate for not knowing:

Java
C#
Go
Rust
TypeScript

when the job says "at least one".

DO NOT put those alternative languages into
missing_important_skills.

Only mark a specific language as missing when the job
explicitly requires that exact language.
3. DATA STRUCTURES AND ALGORITHMS
Treat these as related evidence:

Data Structures
DSA
Data Structures and Algorithms
Algorithms
Algorithm Development
Problem Solving
Competitive Programming
LeetCode
Coding Problems

If the resume clearly demonstrates DSA or algorithmic
problem solving, do not incorrectly mark algorithms as missing.

Do not invent experience that is not present.
4. OBJECT-ORIENTED PROGRAMMING
Treat these as related:

OOP
OOPS
Object-Oriented Programming
Object-Oriented Design
Classes
Objects
Inheritance
Polymorphism
Encapsulation
Abstraction

If the resume clearly demonstrates OOP,
consider the OOP/design requirement satisfied.
5. EDUCATION

Match the candidate's education against the job requirements.

Treat equivalent wording as equivalent.

Examples:

B.Tech CSE
Computer Science Engineering
Computer Science and Engineering
Bachelor's in Computer Science

should generally be considered equivalent when appropriate.

Do not penalize wording differences.
6. INTERNSHIP AND PROJECT EXPERIENCE
Consider:

Technical internships
Software internships
Academic projects
Personal software projects
Full-stack projects
AI/ML projects
Backend projects
Frontend projects
Open-source projects

Project experience is valid evidence of technical ability
when the job does not require formal professional experience.

7. PREFERRED SKILLS
Preferred skills should improve the score,
but should NOT dominate the score.

Examples:
AWS
Cloud
SQL
NoSQL
Git
GitHub
AI tools
Open-source
Debugging
SDLC
CI/CD
Docker
Linux
Agile
Scrum
Problem solving
Communication

8. MATCHING SKILLS
matching_skills should contain skills clearly supported
by the candidate's resume.

Do not invent skills.

Use semantic matching.

For example:

OOP and OOPS are equivalent.

Git and GitHub can support version-control evidence.

Data Structures and DSA are related.

9. MISSING IMPORTANT SKILLS
Only list genuinely important requirements that are not
demonstrated in the resume.

Do NOT list every optional technology.

Do NOT list alternative programming languages when
"at least one" language has already been satisfied.

Do not unfairly penalize candidates for optional skills.

10. EXPERIENCE REQUIREMENT

If no minimum experience is specified:

experience_requirement_met = null

If a minimum experience requirement exists:

true only if clearly satisfied.

Do not invent experience.

11. AGE
If the job contains an age requirement but the resume
does not provide age information:

Do NOT reduce the technical score because age cannot
be verified from the resume.
12. EDUCATION MATCH
education_match should be:

true

when the resume clearly satisfies the education requirement.

Otherwise:

false

If education information is genuinely unavailable,
use null.

13. FINAL VERDICT
Write a short recruiter-style verdict.

Mention:

- Major strengths
- Important gaps
- Overall fit

Do not reject a candidate merely because they do not
know every programming language listed as alternatives.
14. CONSISTENCY
Use the SAME evaluation logic for every candidate.

Do not randomly change scoring criteria between candidates.

A candidate with stronger evidence of required skills
should generally receive a higher score.

15. DO NOT INVENT

Only use information actually present in the structured resume.

Do not assume AWS, Java, Docker, Kubernetes, etc.
unless the resume demonstrates them.


Return ONLY JSON.
"""

    messages = [
        {
            "role": "system",
            "content": (
                "You are a consistent technical recruiting "
                "evaluation engine. Follow the evaluation "
                "rules exactly."
            ),
        },
        {
            "role": "user",
            "content": prompt,
        },
    ]

    raw_output = call_llm(messages)

    try:

        data = json.loads(
            raw_output
        )

        result = MatchResult(
            **data
        )

        return result

    except (
        json.JSONDecodeError,
        ValidationError
    ) as e:

        print(
            "Matching error:",
            e
        )

        print(
            "Raw output:"
        )

        print(
            raw_output
        )

        raise
# 11. PRINT CANDIDATE DETAILS
def print_candidate_details(
    candidate,
    rank
):

    details = candidate["details"]

    print(
        f"\n{rank}. {candidate['name']}"
    )

    print(
        f"Score: {candidate['score']:.2f}%"
    )

    print(
        "Matching Skills:",
        ", ".join(
            details.get(
                "matching_skills",
                []
            )
        )
        or "None"
    )

    print(
        "Missing Important Skills:",
        ", ".join(
            details.get(
                "missing_important_skills",
                []
            )
        )
        or "None"
    )

    print(
        "Experience Requirement Met:",
        details.get(
            "experience_requirement_met"
        )
    )

    print(
        "Education Match:",
        details.get(
            "education_match"
        )
    )

    print(
        "Verdict:",
        details.get(
            "final_verdict",
            ""
        )
    )
# 12. MAIN
def main():

    print("=" * 70)
    print(
        "AI RESUME SCREENING SYSTEM"
    )
    print("=" * 70)

    print(
        "\nAnalyzing Job Description..."
    )
    # Parse Job
    job = parse_job_description(
        job_description
    )

    print(
        "\nJOB INFORMATION"
    )

    print(
        "-" * 70
    )

    print(
        "Role:",
        job.role
    )

    print(
        "Required Skills:",
        ", ".join(
            job.required_skills
        )
        if job.required_skills
        else "None"
    )

    print(
        "Preferred Skills:",
        ", ".join(
            job.preferred_skills
        )
        if job.preferred_skills
        else "None"
    )

    print(
        "Minimum Experience:",
        job.minimum_experience
        if job.minimum_experience is not None
        else "Not specified"
    )

    print(
        "Education:",
        ", ".join(
            job.education_requirements
        )
        if job.education_requirements
        else "None"
    )
    # Resume Folder
    resume_folder = (
        Path(__file__).parent
        / "resumes"
    )

    if not resume_folder.exists():

        raise FileNotFoundError(
            f"Resume folder nahi mila:\n"
            f"{resume_folder}"
        )
    # Find resumes
    resume_files = [
        file_path
        for file_path in resume_folder.iterdir()
        if file_path.is_file()
        and file_path.suffix.lower()
        in [".pdf", ".docx"]
    ]

    if not resume_files:

        print(
            "\nNo PDF/DOCX resumes found."
        )

        return
    # Sort files for consistent processing order
    resume_files.sort(
        key=lambda x: x.name.lower()
    )

    print("\n")

    print("=" * 70)

    print(
        f"FOUND {len(resume_files)} RESUMES"
    )

    print("=" * 70)

    all_results = []
    # PROCESS EACH RESUME
    for index, file_path in enumerate(
        resume_files,
        start=1
    ):

        print("\n")

        print("=" * 70)

        print(
            f"PROCESSING "
            f"{index}/{len(resume_files)}"
        )

        print(
            file_path.name
        )

        print("=" * 70)

        try:
            # Extract text
            resume_text = read_resume(
                file_path
            )

            if (
                not resume_text
                or not resume_text.strip()
            ):

                print(
                    "Could not extract text "
                    "from resume."
                )

                continue

            print(
                "Resume text extracted."
            )
            # Parse resume
            print(
                "Parsing resume with AI..."
            )

            parsed_resume = parse_resume(
                resume_text
            )

            print(
                "Candidate:",
                parsed_resume.name
                or file_path.stem
            )
            # Score
            print(
                "Calculating match score..."
            )

            result = final_score(
                job,
                parsed_resume
            )

            print(
                f"Score: "
                f"{result.score:.2f}%"
            )
            # Store result
            all_results.append(
                {
                    "name": (
                        parsed_resume.name
                        or file_path.stem
                    ),
                    "score": result.score,
                    "details": (
                        result.details.model_dump()
                    ),
                }
            )
            # Small delay
            if index < len(resume_files):

                time.sleep(2)

        except Exception as e:

            print(
                f"FAILED: "
                f"{file_path.name}"
            )

            print(
                f"Reason: {e}"
            )

            continue
    # NO RESULTS
    if not all_results:

        print(
            "\nNo resumes were successfully "
            "processed."
        )

        return
    # FINAL RANKING
    all_results.sort(
        key=lambda candidate:
        candidate["score"],
        reverse=True
    )

    print("\n\n")

    print("=" * 70)

    print(
        "FINAL RANKING"
    )

    print("=" * 70)

    for rank, candidate in enumerate(
        all_results,
        start=1
    ):

        print(
            f"{rank}. "
            f"{candidate['name']} "
            f"- "
            f"{candidate['score']:.2f}%"
        )
    # TOP 2
    top_2 = all_results[:2]

    print("\n\n")

    print("=" * 70)

    print(
        "TOP 2 CANDIDATES"
    )

    print("=" * 70)

    for rank, candidate in enumerate(
        top_2,
        start=1
    ):

        print_candidate_details(
            candidate,
            rank
        )
    # LOWEST 2
    worst_2 = (
        all_results[-2:][::-1]
    )

    print("\n\n")

    print("=" * 70)

    print(
        "LOWEST 2 CANDIDATES"
    )

    print("=" * 70)

    for rank, candidate in enumerate(
        worst_2,
        start=1
    ):

        print_candidate_details(
            candidate,
            rank
        )
    # COMPLETED
    print("\n")

    print("=" * 70)

    print(
        "SCREENING COMPLETED"
    )

    print("=" * 70)
# 13. RUN PROGRAM

if __name__ == "__main__":
    main()
